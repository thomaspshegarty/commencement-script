import xlrd
import classes
from classes import Student
from classes import Major
from classes import Degree
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#column 0: first name
#column 1: last name
#column 2: pronunciation
#column 3: email address
#column 4: degree level
#column 5: major 1
#column 6: major 2
#column 7: minor 1
#column 8: minor 2
#column 9: special programs
#column 10: cluster living
#column 11: graduate degrees 1
#column 12: graduate degrees other
#column 13: graduate degrees 1
#column 14: graduate degrees other
#column 15: study abroad
#column 16: honors and scholarships
#column 17: extracurricular
#column 18: dissertation
#column 19: post-graduation plans
#column 20: special message	

#definitions

def construct_student(worksheet, i):
	new_student = Student(worksheet, i)
	return new_student;

def students_to_degrees(student_list = []):
	degree_list = list()

	for s in student_list:
		student_degree = s.degree_level
		degree_exists = False
		degree = Degree()
		for d in degree_list:
			if d.degree_name == student_degree:
				degree_exists = True
				degree = d;
		if degree_exists:
			degree.add_student(s);
		else:
			degree.set_name(student_degree)
			degree.add_student(s)
			degree_list.append(degree);
	return degree_list;

def write_student(document, student):
	p = document.add_paragraph()
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	first_name = student.first_name.decode('utf-8').encode('ascii','ignore')
	last_name = student.last_name.decode('utf-8').encode('ascii','ignore')
	pronunciation = student.pronunciation.decode('utf-8').encode('ascii','ignore')
	degree = student.degree_level.decode('utf-8').encode('ascii','ignore')
	major = classes.get_major(student.degree_level,student)
	minor = classes.get_minor(student)
	study_abroad = student.study_abroad.decode('utf-8').encode('ascii','ignore').rstrip()
	honors = student.honors.decode('utf-8').encode('ascii','ignore').rstrip()
	extra = student.extra.decode('utf-8').encode('ascii','ignore').rstrip()
	post = student.post_grad.decode('utf-8').encode('ascii','ignore').rstrip()

	p.add_run(first_name +" "+last_name).font.size = Pt(20)
	p.add_run("\nPronunciation: " + pronunciation)
	p.add_run("\nDegree: " + degree+", ")
	p.add_run("Major: "+major)
	if(minor != ''):
		p.add_run(" Minor: "+minor)
	p.add_run("\nStudy Abroad Experience: "+study_abroad)
	p.add_run("\nHonors and Scholarships: "+honors)
	p.add_run("\nExtracurricular Activities: "+extra)
	if(student.degree_level == 'PhD'):
		p.add_run("\nDissertation: "+student.dissertation.decode('utf-8').encode('ascii','ignore'));
	p.add_run("\nPost-graduation Plans: "+post)
	p.add_run("\n")

#script

workbook = xlrd.open_workbook('commencement_part_2.xls')
worksheet = workbook.sheet_by_index(0)

document = Document()
document.add_heading('Commencement Cards',0)

students_list = list()

i = 1
while i < 95:
	student = construct_student(worksheet, i)
	students_list.append(student)
	i=i+1;

degree_list = students_to_degrees(students_list)

for d in degree_list:
	d.order_majors()

degree_list = sorted(degree_list, key = lambda degree: degree.degree_name)
for d in degree_list:
	d.majors = sorted(d.majors, key = lambda major: major.major_name);
	for m in d.majors:
		m.students = sorted(m.students, key = lambda student: student.last_name);

for d in degree_list:
	for m in d.majors:
		for s in m.students:
			print s.last_name
			write_student(document, s);

document.save('card_data.docx')


